#!/usr/bin/env python3
"""TIG (teljesítésigazolás) .docx generálás - általános és ERSTE típus.

Kétféle TIG-et állít elő, mert a cég kétfélét használ:

- "altalanos": a legtöbb ügyfélnél ezt használjuk. Egyszerű felépítés, kódból
  épül fel, így rugalmas (AAM vs. +Áfa, tetszőleges igazoló cég, több soros
  megrendelés-tárgy).
- "erste": az ERSTE-s bevásárlóközpontok (NFP / TC / CK) havi TIG-je. Ez az
  ügyfél saját, fejléces formája, ezért nem újraépítjük, hanem az
  assets/tig_erste_sablon.docx placeholdereit töltjük ki - így a logó, a
  margók és az aláírótábla pontosan olyan marad, amit az ügyfél megszokott.

Az összeg betűs alakját mindig a szamnev modul adja, sosem kézzel írjuk be.

Bemenet: JSON fájl (több TIG is lehet benne egyszerre, mert a havi zárásnál
általában 3-4 készül egyszerre). Példa a references/tig.md fájlban.

    python3 tig_general.py bemenet.json --kimenet ./out
    python3 tig_general.py --minta            # mintabemenet kiírása
"""
import argparse
import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from szamnev import szamnev, ezres  # noqa: E402

try:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    sys.exit("Hiányzik a python-docx. Telepítsd: pip install python-docx")

SKILL_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ERSTE_SABLON = os.path.join(SKILL_DIR, "assets", "tig_erste_sablon.docx")
UGYFELEK = os.path.join(SKILL_DIR, "references", "ugyfelek.json")


# ---------------------------------------------------------------- általános

def _cimke_ertek_tabla(doc, sorok):
    """Kétoszlopos címke/érték tábla, ahogy az általános TIG-en áll."""
    t = doc.add_table(rows=len(sorok), cols=2)
    t.style = "Table Grid"
    for i, (cimke, ertek) in enumerate(sorok):
        bal = t.rows[i].cells[0]
        jobb = t.rows[i].cells[1]
        bal.text = ""
        jobb.text = ""
        r = bal.paragraphs[0].add_run(cimke)
        r.bold = True
        # A megrendelés tárgya több soros is lehet - a sortöréseket megtartjuk.
        sorok_ertek = str(ertek).split("\n")
        jobb.paragraphs[0].add_run(sorok_ertek[0])
        for extra in sorok_ertek[1:]:
            jobb.add_paragraph(extra)
    return t


def epit_altalanos(adat, kimenet_dir):
    doc = docx.Document()

    # Alap betűtípus - a meglévő TIG-ek talpatlan, 11 pt-os szöveget használnak.
    stilus = doc.styles["Normal"]
    stilus.font.name = "Calibri"
    stilus.font.size = Pt(11)

    cim = doc.add_paragraph()
    cim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cim.add_run("TELJESÍTÉS IGAZOLÁS")
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph()

    _cimke_ertek_tabla(doc, [
        ("Szerződő partner megnevezése:", adat["szerzodo_partner"]),
        ("Megrendelés dátuma:", adat["megrendeles_datuma"]),
        ("A megrendelés tárgya:", adat["megrendeles_targya"]),
    ])

    doc.add_paragraph()
    doc.add_paragraph(
        "A fenti szerződésben/megrendelésben foglalt feladatok teljesítését igazolom."
    )
    doc.add_paragraph("A számla benyújtható az alábbiak szerint:")
    doc.add_paragraph()

    osszeg = adat["osszeg"]
    # "+ Áfa" a fő eset, "AAM" az alanyi adómentes partnereknél (pl. EV-k).
    afa = adat.get("afa", "+ Áfa")
    osszeg_szoveg = f"{ezres(osszeg)} Ft {afa}" if afa != "AAM" else f"{ezres(osszeg)} Ft (AAM)"

    _cimke_ertek_tabla(doc, [
        ("Összege:", osszeg_szoveg),
        ("Fizetés módja:", adat.get("fizetes_modja", "Átutalás")),
        ("Fizetés határideje:", adat["fizetesi_hatarido"]),
    ])

    doc.add_paragraph()
    doc.add_paragraph(f"{adat.get('kelt_hely', 'Szeged')}, {adat['kelt_datum']}")
    doc.add_paragraph()
    doc.add_paragraph("…………………………..")
    doc.add_paragraph(adat["igazolo"])

    ut = os.path.join(kimenet_dir, adat["fajlnev"])
    doc.save(ut)
    return ut


# -------------------------------------------------------------------- ERSTE

def _csere_bekezdesben(par, mezok):
    """{{PLACEHOLDER}} csere futásonként, a formázás megtartásával.

    A placeholderek szándékosan egy-egy futáson belül vannak a sablonban,
    ezért elég futásonként cserélni - nem kell a bekezdést újraépíteni.
    """
    for run in par.runs:
        if "{{" not in run.text:
            continue
        szoveg = run.text
        for kulcs, ertek in mezok.items():
            szoveg = szoveg.replace("{{" + kulcs + "}}", str(ertek))
        run.text = szoveg


def epit_erste(adat, kimenet_dir):
    if not os.path.exists(ERSTE_SABLON):
        sys.exit(f"Nincs meg az ERSTE sablon: {ERSTE_SABLON}")
    doc = docx.Document(ERSTE_SABLON)

    osszeg = adat["osszeg"]
    mezok = {
        "MUNKA_MEGJELOLESE": adat["munka_megjelolese"],
        "TELJESITESI_IDOSZAK": adat["teljesitesi_idoszak"],
        "OSSZEG": ezres(osszeg),
        "OSSZEG_BETUVEL": szamnev(osszeg),
        "KELT_HELY": adat.get("kelt_hely", "Budapest"),
        "KELT_DATUM": adat["kelt_datum"],
    }
    for par in doc.paragraphs:
        _csere_bekezdesben(par, mezok)

    ut = os.path.join(kimenet_dir, adat["fajlnev"])
    doc.save(ut)

    # Biztonsági háló: ha bármi placeholder bennmaradt, az néma hiba lenne.
    ellenoriz = docx.Document(ut)
    maradek = [p.text for p in ellenoriz.paragraphs if "{{" in p.text]
    if maradek:
        sys.exit(f"HIBA: kitöltetlen mező maradt a(z) {adat['fajlnev']} fájlban:\n  "
                 + "\n  ".join(maradek))
    return ut


# --------------------------------------------------------------------- futás

def ugyfel_adatok(kod):
    """Ügyfélkódhoz (pl. 'NFP') tartozó állandó adatok a regiszterből."""
    if not os.path.exists(UGYFELEK):
        return {}
    with open(UGYFELEK, encoding="utf-8") as f:
        return json.load(f).get(kod, {})


MINTA = {
    "tigek": [
        {
            "tipus": "erste",
            "ugyfel": "NFP",
            "teljesitesi_idoszak": "2026.08.01. – 2026.08.31.",
            "osszeg": 809985,
            "kelt_datum": "2026.09.01.",
            "fajlnev": "08_NFP_marketing_TIG_2026.docx"
        },
        {
            "tipus": "altalanos",
            "szerzodo_partner": "Marketing Store Kft.",
            "megrendeles_datuma": "2025. december 19.",
            "megrendeles_targya": "Online marketing kampány menedzsment\nProjektmenedzsment 2026.08. hó",
            "osszeg": 738000,
            "afa": "+ Áfa",
            "fizetes_modja": "Átutalás",
            "fizetesi_hatarido": "2026.09.20.",
            "igazolo": "Arcideál Kft.",
            "kelt_hely": "Szeged",
            "kelt_datum": "2026.09.01.",
            "fajlnev": "Arcideál TIG - Havidíj 2026.08. hó.docx"
        }
    ]
}


def main():
    ap = argparse.ArgumentParser(description="TIG .docx generálás")
    ap.add_argument("bemenet", nargs="?", help="JSON bemeneti fájl")
    ap.add_argument("--kimenet", default=".", help="kimeneti könyvtár")
    ap.add_argument("--minta", action="store_true", help="mintabemenet kiírása")
    args = ap.parse_args()

    if args.minta:
        print(json.dumps(MINTA, ensure_ascii=False, indent=2))
        return
    if not args.bemenet:
        ap.error("adj meg bemeneti JSON fájlt (vagy --minta)")

    with open(args.bemenet, encoding="utf-8") as f:
        adatok = json.load(f)
    os.makedirs(args.kimenet, exist_ok=True)

    for tig in adatok["tigek"]:
        tipus = tig.get("tipus", "altalanos")
        if tipus == "erste":
            # Az ügyfélkódból kipótoljuk az állandó mezőket (munka megjelölése),
            # hogy a szerződés szövegét ne kelljen havonta újragépelni.
            if "munka_megjelolese" not in tig and tig.get("ugyfel"):
                allando = ugyfel_adatok(tig["ugyfel"])
                if "munka_megjelolese" in allando:
                    tig["munka_megjelolese"] = allando["munka_megjelolese"]
                tig.setdefault("kelt_hely", allando.get("kelt_hely", "Budapest"))
            ut = epit_erste(tig, args.kimenet)
        elif tipus == "altalanos":
            ut = epit_altalanos(tig, args.kimenet)
        else:
            sys.exit(f"Ismeretlen TIG típus: {tipus!r}")
        osszeg = tig["osszeg"]
        print(f"kész: {ut}")
        print(f"      {ezres(osszeg)} Ft — azaz {szamnev(osszeg)} forint")


if __name__ == "__main__":
    main()
